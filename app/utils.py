from docx import Document
from datetime import datetime
import re
from app import mongo
from docx.oxml.ns import qn

def clean_text(text):
    """Clean up Word-specific math characters and convert common symbols to LaTeX equivalents"""
    if not text:
        return ""
    
    # 1. Remove Word's "garbage" characters (linearized math indicators)
    # \u2592 = ▒, \u2591 = ░, \u2593 = ▓
    text = text.replace('\u2592', '').replace('\u2591', '').replace('\u2593', '')
    
    # 2. Replace Word's grouping brackets with LaTeX braces
    # \u3016 = 〖, \u3017 = 〗
    text = text.replace('\u3016', '{').replace('\u3017', '}')
    
    # 3. Handle combining marks (e.g., X̄ -> \bar{X}, X̂ -> \hat{X})
    text = re.sub(r'([a-zA-Z0-9])\u0304', r'\\bar{\1}', text) # Combining macron
    text = re.sub(r'([a-zA-Z0-9])\u0302', r'\\hat{\1}', text)  # Combining circumflex
    
    # 4. Convert common Unicode math symbols to LaTeX equivalents
    replacements = {
        '∑': r'\sum ',
        '√': r'\sqrt',
        '∫': r'\int ',
        '∆': r'\Delta ',
        'δ': r'\delta ',
        'α': r'\alpha ',
        'β': r'\beta ',
        'γ': r'\gamma ',
        'π': r'\pi ',
        '∞': r'\infty ',
        '±': r'\pm ',
        '×': r'\times ',
        '÷': r'\div ',
        '≈': r'\approx ',
        '≠': r'\neq ',
        '≤': r'\le ',
        '≥': r'\ge ',
        '→': r'\to ',
        'λ': r'\lambda ',
        'σ': r'\sigma ',
        'μ': r'\mu ',
        'η': r'\eta ',
        'ρ': r'\rho ',
        'θ': r'\theta ',
        'φ': r'\phi ',
        'ω': r'\omega ',
        'Ω': r'\Omega ',
        '^': r'^',
        '_': r'_',
    }
    
    for char, latex in replacements.items():
        text = text.replace(char, latex)
        
    return text

def get_paragraph_text(paragraph):
    """Extract text from a paragraph, including math elements (OMML) with cleanup and LaTeX delimiters"""
    text_parts = []
    
    # Iterate through runs and math blocks in order
    # Using xpath to get children elements in sequence
    for child in paragraph._element.xpath('./w:r | ./m:oMath | ./m:oMathPara | ./w:hyperlink | ./w:ins'):
        if child.tag in (qn('m:oMath'), qn('m:oMathPara')):
            # This is a math block
            m_parts = []
            for t in child.iter(qn('m:t')):
                if t.text:
                    m_parts.append(t.text)
            
            math_content = "".join(m_parts)
            cleaned_math = clean_text(math_content)
            
            if cleaned_math.strip():
                # Wrap math in MathJax delimiters
                text_parts.append(f" ${cleaned_math.strip()}$ ")
        else:
            # Regular text run
            r_parts = []
            for t in child.iter(qn('w:t')):
                if t.text:
                    r_parts.append(t.text)
            
            # Even regular text might contain math symbols if copy-pasted or auto-converted
            # so we clean it too, but don't wrap in $$ unless it was already math
            text_parts.append(clean_text("".join(r_parts)))
    
    # Fallback to standard text if custom extraction yielded nothing
    if not "".join(text_parts).strip() and paragraph.text:
        return clean_text(paragraph.text)
        
    return "".join(text_parts)

def import_from_docx(file, subject_id=None):
    """Import questions from Word document"""
    try:
        doc = Document(file)
    except Exception as e:
        raise Exception(f"Không thể đọc file Word: {str(e)}")
    
    questions = []
    current_question = None
    
    for i, paragraph in enumerate(doc.paragraphs):
        text = get_paragraph_text(paragraph).strip()
        
        if not text:
            continue
        
        # Detect question (starts with number followed by . or ) or a period, or starts with "Câu")
        question_match = re.match(r'^(\d+)[\.)\]]\s+(.*)', text) or re.match(r'^Câu\s+(\d+)[\.)\:\-\s]+(.*)', text, re.IGNORECASE)
        
        if question_match:
            if current_question and current_question['question'] and current_question['correct_answer']:
                questions.append(current_question)
            
            question_text = question_match.group(2) if question_match.groups() > 1 else text
            
            current_question = {
                'question': question_text.strip(),
                'options': {},
                'correct_answer': None,
                'category': 'Thủy văn công trình',
                'difficulty': 'medium'
            }
        
        # Detect options (starts with a), b), c), d) or A., B., etc.)
        elif re.match(r'^([a-dA-D])[\.)\]]\s+(.*)', text) or re.match(r'^([a-dA-D])\)\s*(.*)', text):
            if current_question:
                match = re.match(r'^([a-dA-D])[\.)\]]\s*(.*)', text) or re.match(r'^([a-dA-D])\)\s*(.*)', text)
                if match:
                    option_key = match.group(1).lower()
                    option_text = match.group(2).strip()
                    current_question['options'][option_key] = option_text
        
        # Detect correct answer (multiple formats)
        elif re.search(r'(Đáp án|Đáp Án|ĐÁP ÁN|Answer|ANSWER)[:\s]+([a-dA-D])', text, re.IGNORECASE):
            if current_question:
                match = re.search(r'(?:Đáp án|Đáp Án|ĐÁP ÁN|Answer|ANSWER)[:\s]+([a-dA-D])', text, re.IGNORECASE)
                if match:
                    current_question['correct_answer'] = match.group(1).lower()
        
        elif re.match(r'^[a-dA-D]$', text.strip(), re.IGNORECASE) and current_question and not current_question['correct_answer']:
            # Sometimes answer is just a single letter on its own line
            current_question['correct_answer'] = text.strip().lower()
    
    # Add the last question
    if current_question and current_question['question'] and current_question['correct_answer']:
        questions.append(current_question)
    
    # Save to database
    count = 0
    for q in questions:
        if q['question'] and q['options'] and q['correct_answer']:
            # Check if question already exists
            existing = mongo.db.questions.find_one({
                'question': q['question'],
                'correct_answer': q['correct_answer']
            })
            
            if not existing:
                mongo.db.questions.insert_one({
                    'question': q['question'],
                    'options': q['options'],
                    'correct_answer': q['correct_answer'],
                    'category': q['category'],
                    'difficulty': q['difficulty'],
                    'subject_id': ObjectId(subject_id) if subject_id else None,
                    'created_at': datetime.utcnow()
                })
                count += 1
    
    return count

def generate_sample_docx():
    """Generate a sample Word document with correct format"""
    from docx import Document
    
    doc = Document()
    
    doc.add_heading('Mẫu câu hỏi Thủy Văn Công Trình', 0)
    doc.add_paragraph('Format chuẩn để import câu hỏi vào hệ thống:')
    
    doc.add_paragraph('1. Câu hỏi số 1 về thủy văn?')
    doc.add_paragraph('a) Đáp án A cho câu hỏi 1')
    doc.add_paragraph('b) Đáp án B cho câu hỏi 1')
    doc.add_paragraph('c) Đáp án C cho câu hỏi 1')
    doc.add_paragraph('d) Đáp án D cho câu hỏi 1')
    doc.add_paragraph('Đáp án: a')
    doc.add_paragraph('')
    
    doc.add_paragraph('2. Câu hỏi số 2 về lưu vực sông?')
    doc.add_paragraph('a) Đáp án A')
    doc.add_paragraph('b) Đáp án B')
    doc.add_paragraph('c) Đáp án C')
    doc.add_paragraph('d) Đáp án D')
    doc.add_paragraph('Đáp án: b')
    
    return doc
